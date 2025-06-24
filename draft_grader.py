import os
import re
import logging
import time # For potential delays
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document as DocxDocument
# from docx.shared import Pt # Not strictly needed for basic prose dump
# from docx.enum.text import WD_ALIGN_PARAGRAPH # Not strictly needed

# --- Configuration ---
INPUT_FOLDER = "input_assessments"
OUTPUT_FOLDER = "output_draft_feedback" # Separate output folder for draft feedback
DRAFT_PROMPT_FILE = "draft_feedback_prompt.txt" # New prompt file
REVIEW_PROMPT_FILE = "feedback_review_prompt.txt"  # Prompt for comparing feedback
LOG_FILE = "draft_grading_process.log"
# Folder containing scenario text files
SCENARIO_FOLDER = "Diagnosis scenarios"
# Set to True if you want to add a delay between API calls (e.g., to respect free tier limits)
USE_API_DELAY = True
API_DELAY_SECONDS = 20 # Delay in seconds (e.g., 20-30 for gemini-1.5-flash)

# Setup basic logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler() # Also print to console
    ]
)

# --- Helper Functions ---

def load_api_key():
    """Loads Gemini API key from .env file."""
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        logging.error("GEMINI_API_KEY not found in .env file.")
        raise ValueError("API Key not configured.")
    return api_key

def get_student_identifier_from_filename(filename):
    """
    Extracts a base identifier from the filename.
    Example: "JohnDoe_Assignment1.docx" -> "JohnDoe_Assignment1"
    """
    return os.path.splitext(filename)[0]


def sanitize_for_filename(text):
    """Sanitize text so it is safe for filenames."""
    if not text:
        return ""
    sanitized = re.sub(r"[^A-Za-z0-9_-]+", "_", text.strip())
    return sanitized.strip("_")

def extract_text_from_docx(doc):
    """Extract text from paragraphs and tables in a DOCX Document."""
    text_parts = []
    for para in doc.paragraphs:
        if para.text:
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        text_parts.append(para.text)
    return "\n".join(text_parts)


def extract_text_from_file(filepath):
    """Extracts text and author metadata from supported files.

    Returns a tuple ``(text, author)`` where ``author`` may be ``None`` if not
    available.
    """
    _, extension = os.path.splitext(filepath)
    text = ""
    doc_author = None
    try:
        if extension.lower() == ".docx":
            doc = DocxDocument(filepath)
            doc_author = doc.core_properties.author or None
            for para in doc.paragraphs:
                text += para.text + "\n"

            text = extract_text_from_docx(doc)
        elif extension.lower() == ".pdf":
            # Lazy import PyPDF2 to avoid error if not installed and not processing PDFs
            try:
                import PyPDF2
            except ImportError:
                logging.error("PyPDF2 library is not installed. Please install it to process PDF files: pip install PyPDF2")
                return None, None

            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                    logging.warning(f"PDF '{filepath}' is encrypted. Attempting to read anyway if default password allows.")
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    try:
                        extracted_page_text = page.extract_text()
                        if extracted_page_text: # Ensure text was actually extracted
                             text += extracted_page_text + "\n"
                    except Exception as page_e:
                        logging.warning(f"Could not extract text from page {page_num+1} of {filepath}: {page_e}")
        else: # Attempt plain text for other files
            logging.info(f"Attempting to read '{filepath}' as plain text.")
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        text = re.sub(r'\s{3,}', '\n\n', text).strip()
        if not text.strip():
            logging.warning(f"No text extracted or file is empty: {filepath}")
            return None, None
        return text, doc_author

    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
        return None, None
    except Exception as e: # Catching general exception from PyPDF2 if it was imported
        if 'PyPDF2' in str(type(e)): # Check if it's a PyPDF2 error
             logging.error(f"Could not read PDF (possibly corrupted or password protected): {filepath} - {e}")
             return None, None
        logging.error(f"Error extracting text from {filepath}: {e}")
        return None, None

def load_draft_prompt_template():
    """Loads the draft feedback prompt template from file."""
    try:
        with open(DRAFT_PROMPT_FILE, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        logging.error(f"Draft prompt file '{DRAFT_PROMPT_FILE}' not found.")
        raise
    except Exception as e:
        logging.error(f"Error reading draft prompt file: {e}")
        raise

def load_review_prompt_template():
    """Loads the feedback review prompt template from file."""
    try:
        with open(REVIEW_PROMPT_FILE, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        logging.error(f"Review prompt file '{REVIEW_PROMPT_FILE}' not found.")
        raise
    except Exception as e:
        logging.error(f"Error reading review prompt file: {e}")
        raise

def get_scenario_mapping():
    """Returns a mapping of scenario name keywords to file paths."""
    mapping = {}
    if not os.path.isdir(SCENARIO_FOLDER):
        logging.warning(f"Scenario folder '{SCENARIO_FOLDER}' not found.")
        return mapping
    for fname in os.listdir(SCENARIO_FOLDER):
        if fname.lower().endswith('.txt'):
            base = os.path.splitext(fname)[0]  # e.g., 'sam_d'
            keywords = base.replace('_', ' ').lower()  # 'sam d'
            mapping[keywords] = os.path.join(SCENARIO_FOLDER, fname)
    return mapping


def detect_scenario(student_text, scenario_map):
    """Attempts to detect which scenario is referenced in the student's text."""
    text_lower = student_text.lower()
    for keywords, path in scenario_map.items():
        if keywords in text_lower:
            return path
        first_name = keywords.split()[0]
        if re.search(rf"\b{re.escape(first_name)}\b", text_lower):
            return path
    return None


def load_text_file(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except Exception as e:
        logging.error(f"Failed to load scenario text '{path}': {e}")
        return None

def construct_full_prompt(student_text, master_prompt_template, scenario_text=None):
    """Inserts student (and scenario) text into the prompt template."""
    combined_text = student_text
    if scenario_text:
        combined_text = (
            "### REFERENCED SCENARIO\n" + scenario_text + "\n\n" +
            "### STUDENT SUBMISSION\n" + student_text
        )

    placeholder = "{{STUDENT_SUBMISSION_TEXT_HERE}}"
    if placeholder not in master_prompt_template:
        logging.error(f"Placeholder '{placeholder}' not found in draft prompt template.")
        return master_prompt_template + "\n\n### STUDENT_SUBMISSION_TEXT_TO_GRADE:\n" + combined_text
    return master_prompt_template.replace(placeholder, combined_text)


def construct_prompt_messages(student_text, master_prompt_template, scenario_text=None):
    """Return a list of prompt messages for multi-part API calls."""
    combined_text = student_text
    if scenario_text:
        combined_text = (
            "### REFERENCED SCENARIO\n" + scenario_text + "\n\n" +
            "### STUDENT SUBMISSION\n" + student_text
        )

    placeholder = "{{STUDENT_SUBMISSION_TEXT_HERE}}"
    if placeholder in master_prompt_template:
        pre_prompt = master_prompt_template.split(placeholder)[0].rstrip()
        return [
            pre_prompt,
            combined_text,
            "Please provide draft feedback based on the rubric and instructions above.",
        ]

    logging.warning(
        f"Placeholder '{placeholder}' not found in draft prompt template; sending unsplit prompt."
    )
    return [master_prompt_template + "\n\n" + combined_text]


def call_gemini_api(prompt, api_key):
    """Calls the Gemini API and returns the response text.

    ``prompt`` may be a single string or a list of prompt parts, allowing large
    submissions to be sent as a multi-turn request.
    """
    genai.configure(api_key=api_key)
    # Using gemini-1.5-flash-latest as it has better free tier quotas
    model_name = 'gemini-1.5-flash-latest'
    logging.info(f"Using Gemini model: {model_name}")
    model = genai.GenerativeModel(model_name)
    
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]
    try:
        logging.info("Sending request to Gemini API...")
        response = model.generate_content(prompt, safety_settings=safety_settings)
        
        if not response.parts:
            if response.prompt_feedback and response.prompt_feedback.block_reason:
                logging.error(f"Gemini API request blocked. Reason: {response.prompt_feedback.block_reason_message}")
                return None
            else:
                logging.error("Gemini API returned an empty response with no parts.")
                return None
        
        ai_response_text = response.text
        logging.info("Received response from Gemini API.")
        return ai_response_text.strip() # Strip any leading/trailing whitespace
    except Exception as e:
        logging.error(f"Gemini API call failed: {e}")
        # Log more details if it's a specific Google API error
        if hasattr(e, 'message'):
            logging.error(f"Google API Error Message: {e.message}")
        return None

def review_feedback(student_text, feedback_text, api_key, review_prompt_template=None):
    """Sends student text and AI feedback to Gemini for accuracy review."""
    if review_prompt_template is None:
        try:
            review_prompt_template = load_review_prompt_template()
        except Exception:
            return None

    prompt = review_prompt_template
    if "{{STUDENT_SUBMISSION_TEXT_HERE}}" in prompt:
        prompt = prompt.replace("{{STUDENT_SUBMISSION_TEXT_HERE}}", student_text)
    else:
        logging.warning("Student submission placeholder missing in review prompt template")
        prompt += f"\n\nSTUDENT SUBMISSION:\n{student_text}"

    if "{{DRAFT_FEEDBACK_TEXT_HERE}}" in prompt:
        prompt = prompt.replace("{{DRAFT_FEEDBACK_TEXT_HERE}}", feedback_text)
    else:
        logging.warning("Draft feedback placeholder missing in review prompt template")
        prompt += f"\n\nAI DRAFT FEEDBACK:\n{feedback_text}"

    return call_gemini_api(prompt, api_key)

def _add_formatted_run(paragraph, text):
    """Adds text to a paragraph, interpreting **bold** markdown syntax."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def save_draft_feedback_to_docx(feedback_prose, output_filepath, student_identifier):
    """Saves the AI-generated prose feedback to a nicely formatted DOCX file."""
    try:
        doc = DocxDocument()
        doc.add_heading(f"Draft Feedback Report for: {student_identifier}", level=1)

        bullet_styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]

        for raw_line in feedback_prose.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                doc.add_paragraph()
                continue

            if line.startswith("**") and line.endswith("**") and len(line) > 4:
                heading_text = line.strip("*")
                doc.add_heading(heading_text, level=2)
                continue

            bullet_match = re.match(r"^(\s*)\*\s+(.*)", line)
            if bullet_match:
                indent, bullet_text = bullet_match.groups()
                level = min(len(indent) // 4, 2)
                para = doc.add_paragraph(style=bullet_styles[level])
                _add_formatted_run(para, bullet_text.strip())
                continue

            para = doc.add_paragraph()
            _add_formatted_run(para, line.strip())

        doc.save(output_filepath)
        logging.info(f"Draft feedback report saved to: {output_filepath}")
    except Exception as e:
        logging.error(f"Failed to create DOCX draft report for {student_identifier}: {e}")

# --- Main Processing Logic ---
def main():
    logging.info("Starting AI Draft Feedback Generator...")
    try:
        api_key = load_api_key()
        draft_prompt_template = load_draft_prompt_template()
        review_prompt_template = load_review_prompt_template()
        scenario_map = get_scenario_mapping()
    except Exception as e:
        logging.critical(f"Initialization failed: {e}")
        return

    if not os.path.exists(INPUT_FOLDER):
        logging.error(f"Input folder '{INPUT_FOLDER}' not found.")
        return
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
        logging.info(f"Created output folder for draft feedback: {OUTPUT_FOLDER}")

    processed_files = 0
    successful_feedback_generations = 0
    
    assessment_files = [f for f in os.listdir(INPUT_FOLDER) if os.path.isfile(os.path.join(INPUT_FOLDER, f))]

    for i, filename in enumerate(assessment_files):
        filepath = os.path.join(INPUT_FOLDER, filename)
        
        logging.info(f"--- Processing file ({i+1}/{len(assessment_files)}): {filename} ---")
        processed_files += 1

        student_identifier = get_student_identifier_from_filename(filename)

        extracted_text, doc_author = extract_text_from_file(filepath)
        if not extracted_text:
            logging.warning(f"Skipping {filename} due to text extraction failure or empty content.")
            continue

        word_count = len(extracted_text.split())
        logging.info(f"Extracted approx. {word_count} words from {filename}.")
        if word_count < 50 : # Arbitrary threshold
             logging.warning(f"Extracted text for {filename} is very short ({word_count} words). Feedback might be limited.")

        scenario_path = detect_scenario(extracted_text, scenario_map)
        scenario_text = load_text_file(scenario_path) if scenario_path else None
        if scenario_text:
            logging.info(
                f"Detected scenario text from '{os.path.basename(scenario_path)}' for {filename}."
            )
        else:
            logging.info("No specific scenario detected; proceeding without scenario context.")

        full_prompt = construct_full_prompt(extracted_text, draft_prompt_template, scenario_text)
        prompt_messages = construct_prompt_messages(extracted_text, draft_prompt_template, scenario_text)
        
        # For debugging, save the full prompt sent to the API
        # prompt_debug_path = os.path.join(OUTPUT_FOLDER, f"{student_identifier}_draft_prompt_sent.txt")
        # with open(prompt_debug_path, "w", encoding="utf-8") as pf:
        #    pf.write(full_prompt)
        # logging.info(f"Full draft prompt saved for debugging: {prompt_debug_path}")

        ai_feedback_prose = call_gemini_api(prompt_messages, api_key)

        if not ai_feedback_prose:
            logging.warning(f"Skipping {filename} due to Gemini API call failure or empty response.")
            continue

        output_filename_base = student_identifier
        if doc_author:
            output_filename_base += f"_{sanitize_for_filename(doc_author)}"
        output_docx_path = os.path.join(OUTPUT_FOLDER, f"{output_filename_base}_draft_feedback.docx")

        save_draft_feedback_to_docx(ai_feedback_prose, output_docx_path, student_identifier)

        review_text = review_feedback(extracted_text, ai_feedback_prose, api_key, review_prompt_template)
        if review_text:
            review_path = os.path.join(OUTPUT_FOLDER, f"{output_filename_base}_feedback_review.txt")
            try:
                with open(review_path, "w", encoding="utf-8") as rf:
                    rf.write(review_text)
                logging.info(f"Feedback review saved to: {review_path}")
            except Exception as e:
                logging.error(f"Failed to save feedback review for {student_identifier}: {e}")
        successful_feedback_generations +=1
        logging.info(f"Successfully generated draft feedback for: {filename}")

        # Add delay if processing multiple files to respect API rate limits
        if USE_API_DELAY and i < len(assessment_files) - 1 : # Don't sleep after the last file
            logging.info(f"Waiting for {API_DELAY_SECONDS} seconds before processing next file...")
            time.sleep(API_DELAY_SECONDS)


    logging.info("--- Draft Feedback Generation Complete ---")
    logging.info(f"Total files found in input folder: {len(assessment_files)}")
    logging.info(f"Files attempted for processing: {processed_files}")
    logging.info(f"Successfully generated draft feedback for: {successful_feedback_generations} files")
    logging.info(f"Draft feedback reports saved in: {OUTPUT_FOLDER}")
    logging.info(f"Log file saved at: {LOG_FILE}")

if __name__ == "__main__":
    # Ensure PyPDF2 is available if PDF processing is expected.
    # This is just a check at the start; the import is handled lazily later.
    try:
        import PyPDF2
    except ImportError:
        logging.warning("PyPDF2 library is not installed. PDF files will not be processed. Run 'pip install PyPDF2' to enable PDF support.")
    
    main()
