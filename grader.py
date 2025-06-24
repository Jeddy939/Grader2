import os
from pathlib import Path
import re
import logging
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document as DocxDocument  # To avoid clash with local 'Document'
import PyPDF2
import yaml  # PyYAML

# --- Configuration ---
INPUT_FOLDER = Path("input_assessments")
OUTPUT_FOLDER = Path("output_feedback")
MASTER_PROMPT_FILE = Path("master_prompt.txt")
LOG_FILE = "grading_process.log"
SUMMARY_FILE = "grading_summary.csv"
GRADE_REVIEW_PROMPT_FILE = Path("grade_review_prompt.txt")
RUBRIC_FILE = Path("rubric.yml")

# Setup basic logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler(),  # Also print to console
    ],
)

# --- Helper Functions ---


def load_api_key():
    """Loads Gemini API key from .env file."""
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        logging.error("GEMINI_API_KEY not found in .env file.")
        raise ValueError("API Key not configured.")
    return api_key


def get_student_name_from_filename(filename):
    """
    Attempts to extract a student name from filename.
    Example: "JohnDoe_Assignment1.docx" -> "JohnDoe"
    """
    name_part = Path(filename).stem
    # Simple heuristic: look for parts separated by common delimiters
    # that might indicate a name. This can be improved.
    potential_name = re.split(r"[_\-\s.]", name_part)[0]
    # Check if it looks like a name (e.g., starts with capital)
    if potential_name and potential_name[0].isupper():
        # Further refine if needed, e.g. look for CamelCase or multiple capitalized words
        return potential_name
    return None


def extract_text_from_docx(doc):
    """Extract text from paragraphs and tables in a DOCX Document."""
    text_parts = []
    for para in doc.paragraphs:
        if para.text:
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        text_parts.append(para.text)
    return "\n".join(text_parts)


def extract_text_from_file(filepath):
    """Extracts text and author metadata from supported files."""
    filepath = Path(filepath)
    extension = filepath.suffix
    text = ""
    doc_author = None
    try:
        if extension.lower() == ".docx":
            doc = DocxDocument(filepath)
            doc_author = doc.core_properties.author or None
            text = extract_text_from_docx(doc)
        elif extension.lower() == ".pdf":
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                    logging.warning(
                        f"PDF '{filepath}' is encrypted. Attempting to read anyway if default password allows."
                    )
                    # You might need to handle decryption if password is known: reader.decrypt('')
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n"
        else:  # Attempt plain text for other files
            logging.info(f"Attempting to read '{filepath}' as plain text.")
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

        text = re.sub(r"\s{3,}", "\n\n", text).strip()
        if not text.strip():
            logging.warning(f"No text extracted or file is empty: {filepath}")
            return None, None
        return text, doc_author

    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
        return None, None
    except PyPDF2.errors.PdfReadError:
        logging.error(
            f"Could not read PDF (possibly corrupted or password protected): {filepath}"
        )
        return None, None
    except Exception as e:
        logging.error(f"Error extracting text from {filepath}: {e}")
        return None, None


def load_master_prompt():
    """Loads the master prompt template from file."""
    try:
        with open(MASTER_PROMPT_FILE, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        logging.error(f"Master prompt file '{MASTER_PROMPT_FILE}' not found.")
        raise
    except Exception as e:
        logging.error(f"Error reading master prompt file: {e}")
        raise


def load_grade_review_prompt_template():
    """Loads the grade review prompt template from file."""
    try:
        with open(GRADE_REVIEW_PROMPT_FILE, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        logging.error(f"Grade review prompt file '{GRADE_REVIEW_PROMPT_FILE}' not found.")
        raise
    except Exception as e:
        logging.error(f"Error reading grade review prompt file: {e}")
        raise


def load_rubric_config():
    """Load grading rubric from YAML configuration."""
    try:
        with open(RUBRIC_FILE, "r", encoding="utf-8") as f:
            return yaml.safe_load(f)
    except FileNotFoundError:
        logging.error(f"Rubric file '{RUBRIC_FILE}' not found.")
        raise
    except Exception as e:
        logging.error(f"Error reading rubric file: {e}")
        raise


def construct_full_prompt(
    student_text, master_prompt_template, *, raise_on_missing=True
):
    """Insert student text into the master prompt template.

    Parameters
    ----------
    student_text : str
        The extracted text of the student's submission.
    master_prompt_template : str
        The prompt template loaded from ``master_prompt.txt``.
    raise_on_missing : bool, optional
        If ``True`` (default), a ``ValueError`` is raised when the required
        placeholder is not present in ``master_prompt_template``. The student
        text is still appended to the returned prompt to avoid an empty
        submission.
    """

    placeholder = "{{STUDENT_SUBMISSION_TEXT_HERE}}"

    if placeholder not in master_prompt_template:
        warning_msg = (
            f"Placeholder '{placeholder}' not found in master prompt template. "
            "Student text will be appended to the end of the prompt."
        )
        logging.warning(warning_msg)

        fallback_prompt = (
            master_prompt_template
            + "\n\n### STUDENT_SUBMISSION_TEXT_TO_GRADE:\n"
            + student_text
        )

        if raise_on_missing:
            # Raising an exception alerts the operator to fix the prompt file
            raise ValueError(warning_msg)

        return fallback_prompt

    return master_prompt_template.replace(placeholder, student_text)


def construct_prompt_messages(student_text, master_prompt_template):
    """Return a list of prompt messages for multi-part API calls."""

    placeholder = "{{STUDENT_SUBMISSION_TEXT_HERE}}"
    if placeholder in master_prompt_template:
        pre_prompt = master_prompt_template.split(placeholder)[0].rstrip()
        return [
            pre_prompt,
            student_text,
            "Please grade the submission above according to the rubric and return the YAML as specified.",
        ]

    logging.warning(
        f"Placeholder '{placeholder}' not found in master prompt template; sending unsplit prompt."
    )
    return [master_prompt_template + "\n\n" + student_text]


def call_gemini_api(prompt, api_key):
    """Calls the Gemini API and returns the response text.

    ``prompt`` may be a single string or a list of prompt parts to be sent as a
    multi-turn request. Splitting large prompts can help the model process long
    student submissions more reliably.
    """
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash-latest")  # Or your preferred model
    # Safety settings can be adjusted if needed
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]
    try:
        logging.info("Sending request to Gemini API...")
        # ``prompt`` can be a string or list of strings for multi-part input
        response = model.generate_content(prompt, safety_settings=safety_settings)
        # Check for empty or blocked responses
        if not response.parts:
            if response.prompt_feedback and response.prompt_feedback.block_reason:
                logging.error(
                    f"Gemini API request blocked. Reason: {response.prompt_feedback.block_reason_message}"
                )
                return None
            else:
                logging.error("Gemini API returned an empty response with no parts.")
                return None

        # Assuming the response is text. If it could be multi-modal, access response.text
        ai_response_text = response.text
        logging.info("Received response from Gemini API.")
        return ai_response_text
    except Exception as e:
        logging.error(f"Gemini API call failed: {e}")
        return None


def parse_gemini_yaml_response(response_text):
    """Parses the YAML response from Gemini."""
    if not response_text:
        return None
    try:
        # LLMs can sometimes add markdown backticks around YAML
        cleaned_response = response_text.strip()
        if cleaned_response.startswith("```yaml"):
            cleaned_response = cleaned_response[7:]
        if cleaned_response.startswith("```"):
            cleaned_response = cleaned_response[3:]
        if cleaned_response.endswith("```"):
            cleaned_response = cleaned_response[:-3]

        cleaned_response = cleaned_response.strip()

        # Sanitize unescaped quotes inside quoted values
        def _sanitize_unescaped_quotes(text: str) -> str:
            sanitized_lines = []
            changed = False
            for line in text.splitlines():
                colon_idx = line.find(":")
                if colon_idx != -1:
                    after = line[colon_idx + 1 :].lstrip()
                    if after.startswith("\"") and after.endswith("\"") and len(after) >= 2:
                        inner = after[1:-1]
                        new_inner = re.sub(r'(?<!\\)\"', r'\\"', inner)
                        if new_inner != inner:
                            line = f"{line[: colon_idx + 1]} \"{new_inner}\""
                            changed = True
                sanitized_lines.append(line)
            if changed:
                logging.info("Sanitized unescaped quotes in Gemini YAML response.")
            return "\n".join(sanitized_lines)

        cleaned_response = _sanitize_unescaped_quotes(cleaned_response)

        parsed_data = yaml.safe_load(cleaned_response)
        if "assistant_reasons" in parsed_data:
            return parsed_data
        logging.error(
            f"Parsed YAML does not have expected structure. Parsed: {parsed_data}"
        )
        return None
    except yaml.YAMLError as e:
        logging.error(
            f"Failed to parse YAML response from Gemini: {e}\nRaw response:\n{response_text}"
        )
        return None
    except Exception as e:
        logging.error(
            f"An unexpected error occurred during YAML parsing: {e}\nRaw response:\n{response_text}"
        )
        return None


def compute_overall_grade(breakdown, grade_bands, total_possible):
    """Compute a letter grade from rubric breakdown points."""
    if not isinstance(breakdown, dict):
        return "N/A"

    total_points = 0
    for item in breakdown.values():
        try:
            total_points += int(item.get("points", 0))
        except Exception:
            continue

    if total_points >= grade_bands.get("A", total_possible):
        return "A"
    if total_points >= grade_bands.get("B", 0):
        return "B"
    if total_points >= grade_bands.get("C", 0):
        return "C"

    if total_points >= total_possible * grade_bands.get("D_ratio", 0):
        return "D"
    if total_points >= total_possible * grade_bands.get("E_ratio", 0):
        return "E"
    return "E"


def calculate_final_grade(bands_data, word_count, rubric_config):
    """Apply rubric rules and compute final grade breakdown."""
    criteria_cfg = rubric_config.get("criteria", {})
    grade_bands = rubric_config.get("grade_bands", {})
    total_possible = rubric_config.get("total_points_possible", 0)

    bands = bands_data.copy()

    # Apply rules deterministically
    for rule in rubric_config.get("rules", []):
        name = rule.get("name")
        try:
            condition = rule.get("condition", "")
            local_vars = {f"{k}_band": bands.get(k) for k in bands}
            local_vars["word_count"] = word_count
            if eval(condition, {}, local_vars):
                if rule.get("action") == "set_band":
                    target = rule.get("target")
                    if target:
                        bands[target] = rule.get("band", bands.get(target))
                elif rule.get("action") == "cap_points":
                    target = rule.get("target")
                    points_cap = rule.get("points")
                    if target and target in bands:
                        bands[target] = min(bands[target], points_cap)
        except Exception:
            logging.warning(f"Failed to evaluate rule '{name}'.")

    breakdown = {}
    total_points = 0
    for cid, cfg in criteria_cfg.items():
        band = int(bands.get(cid, 1))
        max_points = int(cfg.get("max_points", band))
        points = min(band, max_points)
        breakdown[cid] = {"band": band, "points": points}
        total_points += points

    overall_grade = compute_overall_grade(breakdown, grade_bands, total_possible)

    return {
        "total_points": total_points,
        "breakdown": breakdown,
        "overall_grade": overall_grade,
    }


def review_grade(student_text, grade_yaml_text, api_key, review_prompt_template=None):
    """Sends student text and the AI's grade to Gemini for fairness review."""
    if review_prompt_template is None:
        try:
            review_prompt_template = load_grade_review_prompt_template()
        except Exception:
            return None

    prompt = review_prompt_template
    if "{{STUDENT_SUBMISSION_TEXT_HERE}}" in prompt:
        prompt = prompt.replace("{{STUDENT_SUBMISSION_TEXT_HERE}}", student_text)
    else:
        logging.warning("Student submission placeholder missing in grade review prompt template")
        prompt += f"\n\nSTUDENT SUBMISSION:\n{student_text}"

    if "{{AI_GRADE_YAML_HERE}}" in prompt:
        prompt = prompt.replace("{{AI_GRADE_YAML_HERE}}", grade_yaml_text)
    else:
        logging.warning("AI grade placeholder missing in grade review prompt template")
        prompt += f"\n\nAI GRADE:\n{grade_yaml_text}"

    return call_gemini_api(prompt, api_key)


def extract_new_grade_from_review(review_text):
    """Attempt to extract a revised overall grade from the review text."""
    if not review_text:
        return None

    # Common patterns such as "grade should be B" or "recommended grade: A"
    patterns = [
        r"grade\s*should\s*be\s*([A-E])",
        r"recommended\s*grade[:\s]+([A-E])",
        r"proposed\s*grade[:\s]+([A-E])",
        r"new\s*grade[:\s]+([A-E])",
        r"should\s*be\s*an?\s*([A-E])",
    ]

    for pat in patterns:
        m = re.search(pat, review_text, re.IGNORECASE)
        if m:
            return m.group(1).upper()
    return None


def apply_criteria_adjustments(parsed_data, adjustments, rubric_config):
    """Apply band changes from review to the parsed YAML data."""
    if not adjustments:
        return

    grade_section = parsed_data.get("assistant_grade", {})
    breakdown = grade_section.get("breakdown", {})
    for crit, new_band in adjustments.items():
        if crit in breakdown:
            breakdown[crit]["band"] = new_band
            max_pts = rubric_config.get("criteria", {}).get(crit, {}).get("max_points", new_band)
            breakdown[crit]["points"] = min(new_band, max_pts)

    total_points = sum(int(item.get("points", 0)) for item in breakdown.values())
    grade_section["total_points"] = total_points
    grade_section["overall_grade"] = compute_overall_grade(
        breakdown,
        rubric_config.get("grade_bands", {}),
        rubric_config.get("total_points_possible", 0),
    )

def extract_criteria_adjustments(review_text):
    """Parse review text for suggested band adjustments using regex."""
    if not review_text:
        return {}

    pattern = re.compile(r"ADJUSTMENT:\s*(\w+)\s*->\s*([1-5])", re.IGNORECASE)
    adjustments = {}
    for match in pattern.finditer(review_text):
        crit = match.group(1)
        band = int(match.group(2))
        adjustments[crit] = band
    return adjustments


def format_feedback_as_docx(
    yaml_data,
    output_filepath,
    student_identifier,
    rubric_config,
    doc_author=None,
    override_grade=None,
):
    """Formats the YAML data into a human-readable DOCX report."""
    try:
        doc = DocxDocument()
        doc.add_heading(f"Feedback Report for: {student_identifier}", level=1)
        if doc_author:
            doc.add_paragraph(f"Author (from file metadata): {doc_author}")

        # Overall Grade and Points
        grade_info = yaml_data.get("assistant_grade", {})
        breakdown = grade_info.get("breakdown", {})

        ai_overall_grade = grade_info.get("overall_grade")
        computed_grade = compute_overall_grade(
            breakdown, rubric_config.get("grade_bands", {}), rubric_config.get("total_points_possible", 0)
        )
        final_grade = override_grade if override_grade else computed_grade
        try:
            total_points = sum(int(item.get("points", 0)) for item in breakdown.values())
        except Exception:
            total_points = grade_info.get("total_points", "N/A")
        max_total_points = rubric_config.get("total_points_possible", 0)

        doc.add_heading("Overall Assessment", level=2)
        if ai_overall_grade:
            doc.add_paragraph(f"AI Reported Grade: {ai_overall_grade}")
        doc.add_paragraph(f"Grade Based on Points: {final_grade}")
        if ai_overall_grade and ai_overall_grade != final_grade:
            doc.add_paragraph("Note: Grade adjusted based on rubric totals.")
        doc.add_paragraph(f"Total Points: {total_points} / {max_total_points}")
        doc.add_paragraph()  # Spacer

        # Criteria Breakdown
        doc.add_heading("Detailed Breakdown by Criterion", level=2)
        reasons = yaml_data.get("assistant_reasons", [])

        rubric_criteria_details = {
            cid: {
                "name": cfg.get("name", cid),
                "max_points": cfg.get("max_points", 0),
            }
            for cid, cfg in rubric_config.get("criteria", {}).items()
        }

        for idx, reason_item in enumerate(reasons, start=1):
            criterion_id = reason_item.get("criterion", "Unknown Criterion")
            band = reason_item.get("band", "N/A")
            rationale = reason_item.get("rationale", "No rationale provided.")
            evidence = reason_item.get("evidence", "No evidence quoted.")

            criterion_details = rubric_criteria_details.get(
                criterion_id,
                {"name": criterion_id.replace("_", " ").title(), "max_points": "N/A"},
            )
            criterion_name = criterion_details["name"]

            criterion_grade_info = breakdown.get(criterion_id, {})
            points_achieved = criterion_grade_info.get("points", "N/A")
            max_criterion_points = criterion_details["max_points"]

            doc.add_heading(f"{idx}. {criterion_name}", level=3)

            doc.add_paragraph(f"Band Achieved: {band}")
            doc.add_paragraph(f"Points: {points_achieved} / {max_criterion_points}")

            doc.add_paragraph("AI's Rationale:", style="Intense Quote")
            doc.add_paragraph(rationale)

            doc.add_paragraph("Evidence from Student's Work:", style="Intense Quote")
            if "\n" in evidence:
                for line in evidence.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(evidence if evidence else "N/A")

            doc.add_paragraph()  # Spacer

        doc.save(output_filepath)
        logging.info(f"Feedback report saved to: {output_filepath}")

    except Exception as e:
        logging.error(f"Failed to create DOCX report for {student_identifier}: {e}")


# --- Main Processing Logic ---
def main():
    logging.info("Starting AI Student Assessment Grader...")
    try:
        api_key = load_api_key()
        master_prompt_template = load_master_prompt()
        rubric_config = load_rubric_config()
    except Exception as e:
        logging.critical(f"Initialization failed: {e}")
        return

    if not INPUT_FOLDER.exists():
        logging.error(f"Input folder '{INPUT_FOLDER}' not found.")
        return
    if not OUTPUT_FOLDER.exists():
        OUTPUT_FOLDER.mkdir(parents=True)
        logging.info(f"Created output folder: {OUTPUT_FOLDER}")

    processed_files = 0
    successful_grades = 0
    summary_entries = []

    for filepath in INPUT_FOLDER.iterdir():
        if not filepath.is_file():
            continue
        filename = filepath.name

        logging.info(f"--- Processing file: {filename} ---")
        processed_files += 1

        student_name_guess = get_student_name_from_filename(filename)
        student_identifier = (
            student_name_guess if student_name_guess else filepath.stem
        )

        extracted_text, doc_author = extract_text_from_file(filepath)
        if not extracted_text:
            logging.warning(
                f"Skipping {filename} due to text extraction failure or empty content."
            )
            continue

        # Simple word count for info, AI will use its own logic based on rubric
        word_count = len(extracted_text.split())
        logging.info(f"Extracted approx. {word_count} words from {filename}.")
        if word_count < 50:  # Arbitrary threshold for very short/empty files
            logging.warning(
                f"Extracted text for {filename} is very short ({word_count} words). May not be suitable for grading."
            )
            # continue # Optional: skip very short files

        full_prompt = construct_full_prompt(extracted_text, master_prompt_template)
        prompt_messages = construct_prompt_messages(extracted_text, master_prompt_template)

        # For debugging, you might want to save the full prompt
        # with open(os.path.join(OUTPUT_FOLDER, f"{student_identifier}_prompt.txt"), "w", encoding="utf-8") as pf:
        #    pf.write(full_prompt)

        api_response = call_gemini_api(prompt_messages, api_key)
        if not api_response:
            logging.warning(f"Skipping {filename} due to Gemini API call failure.")
            continue

        parsed_data = parse_gemini_yaml_response(api_response)
        if not parsed_data:
            logging.warning(f"Skipping {filename} due to YAML parsing failure.")
            # Save raw response for debugging
            raw_response_path = OUTPUT_FOLDER / f"{student_identifier}_raw_gemini_response.txt"
            with open(raw_response_path, "w", encoding="utf-8") as f:
                f.write(api_response if api_response else "No response received.")
            logging.info(f"Raw Gemini response saved to: {raw_response_path}")
            continue

        # Calculate grade using rubric
        bands = {
            item.get("criterion"): int(item.get("band", 1))
            for item in parsed_data.get("assistant_reasons", [])
            if item.get("criterion")
        }
        parsed_data["assistant_grade"] = calculate_final_grade(
            bands, word_count, rubric_config
        )

        output_filename_base = student_identifier
        output_docx_path = OUTPUT_FOLDER / f"{output_filename_base}_graded.docx"

        review_text = review_grade(extracted_text, api_response, api_key)
        override_grade = None
        if review_text:
            override_grade = extract_new_grade_from_review(review_text)
            review_path = OUTPUT_FOLDER / f"{output_filename_base}_grade_review.txt"
            try:
                with open(review_path, "w", encoding="utf-8") as rf:
                    rf.write(review_text)
                logging.info(f"Grade review saved to: {review_path}")
                if override_grade:
                    logging.info(
                        f"Applying grade override from review: {override_grade}"
                    )
                adjustments = extract_criteria_adjustments(review_text)
                if adjustments:
                    apply_criteria_adjustments(parsed_data, adjustments, rubric_config)
                    logging.info(f"Applied criterion adjustments: {adjustments}")
            except Exception as e:
                logging.error(f"Failed to save grade review for {student_identifier}: {e}")

        breakdown = parsed_data.get("assistant_grade", {}).get("breakdown", {})
        try:
            total_points = sum(int(item.get("points", 0)) for item in breakdown.values())
        except Exception:
            total_points = parsed_data.get("assistant_grade", {}).get("total_points", "N/A")
        overall_grade = compute_overall_grade(
            breakdown,
            rubric_config.get("grade_bands", {}),
            rubric_config.get("total_points_possible", 0),
        )
        if override_grade:
            overall_grade = override_grade

        summary_entries.append((student_identifier, total_points, overall_grade))

        format_feedback_as_docx(
            parsed_data,
            output_docx_path,
            student_identifier,
            rubric_config,
            doc_author=doc_author,
            override_grade=override_grade,
        )
        successful_grades += 1
        logging.info(f"Successfully processed and graded: {filename}")

    logging.info("--- Processing Complete ---")
    logging.info(
        f"Total files found: {len(list(INPUT_FOLDER.iterdir()))}"
    )  # This will count folders too, refine if needed
    logging.info(f"Files attempted for processing: {processed_files}")
    logging.info(f"Successfully graded: {successful_grades}")
    logging.info(f"Reports saved in: {OUTPUT_FOLDER}")
    logging.info(f"Log file saved at: {LOG_FILE}")

    if summary_entries:
        summary_path = OUTPUT_FOLDER / SUMMARY_FILE
        try:
            with open(summary_path, "w", encoding="utf-8") as sf:
                sf.write("student,total_points,grade\n")
                for ident, points, grade in summary_entries:
                    sf.write(f"{ident},{points},{grade}\n")
            logging.info(f"Summary saved to: {summary_path}")
        except Exception as e:
            logging.error(f"Failed to write summary file: {e}")


if __name__ == "__main__":
    main()
